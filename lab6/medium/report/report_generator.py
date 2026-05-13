from docx import Document
from openpyxl import Workbook

def save_doc(results):
    d = Document()
    d.add_heading("Отчёт по расчётам", level=1)

    for key, value in results.items():
        d.add_paragraph(f"{key}: {value}")

    d.save("report.docx")

def save_xls(results):
    e = Workbook()
    es = e.active

    es.append(["Параметр", "Значение"])
    for key, value in results.items():
        es.append([key, value])

    e.save("report.xlsx")
