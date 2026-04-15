from docx import Document
from openpyxl import Workbook


def save_doc(shape):
    doc = Document()

    doc.add_heading('Отчёт по фигуре', 0)

    doc.add_paragraph(str(shape))
    doc.add_paragraph(f"Площадь: {shape.s()}")
    doc.add_paragraph(f"R описанной: {safe(shape.opis_radius())}")
    doc.add_paragraph(f"r вписанной: {safe(shape.vpis_radius())}")

    doc.save("report.docx")


def save_xls(shape):
    wb = Workbook()
    ws = wb.active

    ws.append(["Фигура", shape.name])
    ws.append(["Площадь", shape.s()])
    ws.append(["R описанной", safe(shape.opis_radius())])
    ws.append(["r вписанной", safe(shape.vpis_radius())])

    wb.save("report.xlsx")


def safe(value):
    return value if value is not None else "нет"
